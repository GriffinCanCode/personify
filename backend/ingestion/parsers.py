import os
from pathlib import Path
from typing import Dict, Optional
from datetime import datetime
import hashlib
import json

# Document parsing
from PyPDF2 import PdfReader
import docx
import markdown

class DocumentParser:
    """Base class for document parsers"""
    
    @staticmethod
    def parse(file_path: str) -> Dict:
        """Parse document and return content with metadata"""
        raise NotImplementedError

class TextParser(DocumentParser):
    """Parser for plain text files"""
    
    @staticmethod
    def parse(file_path: str) -> Dict:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        return {
            'content': content,
            'metadata': {
                'filename': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'modified_date': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
            }
        }

class PDFParser(DocumentParser):
    """Parser for PDF files"""
    
    @staticmethod
    def parse(file_path: str) -> Dict:
        reader = PdfReader(file_path)
        content = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                content.append(text)
        
        full_content = '\n\n'.join(content)
        
        return {
            'content': full_content,
            'metadata': {
                'filename': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'num_pages': len(reader.pages),
                'modified_date': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
            }
        }

class DocxParser(DocumentParser):
    """Parser for Word documents"""
    
    @staticmethod
    def parse(file_path: str) -> Dict:
        doc = docx.Document(file_path)
        
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        full_content = '\n\n'.join(content)
        
        return {
            'content': full_content,
            'metadata': {
                'filename': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'num_paragraphs': len(doc.paragraphs),
                'modified_date': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
            }
        }

class MarkdownParser(DocumentParser):
    """Parser for Markdown files"""
    
    @staticmethod
    def parse(file_path: str) -> Dict:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            md_content = f.read()
        
        # Convert to plain text (strip markdown syntax)
        html = markdown.markdown(md_content)
        # Simple HTML tag removal (for basic parsing)
        import re
        text = re.sub('<[^<]+?>', '', html)
        
        return {
            'content': text,
            'raw_markdown': md_content,
            'metadata': {
                'filename': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'modified_date': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
            }
        }

class JSONParser(DocumentParser):
    """Parser for JSON files (e.g., exported chats, data dumps)"""
    
    @staticmethod
    def parse(file_path: str) -> Dict:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Extract text content from JSON
        content = JSONParser._extract_text(data)
        
        return {
            'content': content,
            'raw_data': data,
            'metadata': {
                'filename': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'modified_date': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
            }
        }
    
    @staticmethod
    def _extract_text(obj, depth=0, max_depth=10) -> str:
        """Recursively extract text from JSON structure"""
        if depth > max_depth:
            return ""
        
        if isinstance(obj, str):
            return obj
        elif isinstance(obj, dict):
            texts = []
            for key, value in obj.items():
                # Prioritize content-like keys
                if key.lower() in ['text', 'content', 'message', 'body', 'description']:
                    texts.insert(0, JSONParser._extract_text(value, depth + 1, max_depth))
                else:
                    texts.append(JSONParser._extract_text(value, depth + 1, max_depth))
            return '\n'.join(filter(None, texts))
        elif isinstance(obj, list):
            return '\n\n'.join(JSONParser._extract_text(item, depth + 1, max_depth) for item in obj)
        else:
            return str(obj)

class AudioParser(DocumentParser):
    """Parser for audio files using Whisper API"""
    
    @staticmethod
    def parse(file_path: str, use_whisper: bool = True) -> Dict:
        if not use_whisper:
            return {
                'content': '',
                'error': 'Audio transcription requires Whisper API',
                'metadata': {
                    'filename': os.path.basename(file_path),
                    'file_size': os.path.getsize(file_path)
                }
            }
        
        from openai import OpenAI
        from backend.config import settings
        
        client = OpenAI(api_key=settings.OPENAI_API_KEY)
        
        with open(file_path, 'rb') as audio_file:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                response_format="verbose_json"
            )
        
        return {
            'content': transcript.text,
            'metadata': {
                'filename': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'duration': transcript.duration if hasattr(transcript, 'duration') else None,
                'language': transcript.language if hasattr(transcript, 'language') else None,
                'modified_date': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
            }
        }

class ParserFactory:
    """Factory to get appropriate parser for file type"""
    
    PARSERS = {
        '.txt': TextParser,
        '.text': TextParser,
        '.pdf': PDFParser,
        '.docx': DocxParser,
        '.doc': DocxParser,
        '.md': MarkdownParser,
        '.markdown': MarkdownParser,
        '.json': JSONParser,
        '.mp3': AudioParser,
        '.m4a': AudioParser,
        '.wav': AudioParser,
        '.mp4': AudioParser,
    }
    
    @staticmethod
    def get_parser(file_path: str) -> Optional[DocumentParser]:
        """Get parser for file based on extension"""
        ext = Path(file_path).suffix.lower()
        return ParserFactory.PARSERS.get(ext)
    
    @staticmethod
    def parse(file_path: str) -> Optional[Dict]:
        """Parse file using appropriate parser"""
        parser = ParserFactory.get_parser(file_path)
        if parser:
            try:
                return parser.parse(file_path)
            except Exception as e:
                return {
                    'content': '',
                    'error': str(e),
                    'metadata': {'filename': os.path.basename(file_path)}
                }
        return None

def compute_content_hash(content: str) -> str:
    """Compute SHA-256 hash of content for deduplication"""
    return hashlib.sha256(content.encode('utf-8')).hexdigest()

